from __future__ import annotations
from pathlib import Path
from typing import Optional, Tuple
from docx import Document as load_docx            # 函式（工廠）
from docx.document import Document as DocxDocument  # 類別（型別用）
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
import zipfile
from io import BytesIO
import lxml.etree as etree

def _set_paragraph_format(paragraph):
    for run in paragraph.runs:
        run.font.name = "新細明體"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
        run.font.size = Pt(12)
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT  # ← 新增：統一置左

def open_docx(path: Path) -> DocxDocument:
    return load_docx(str(path))

def save_docx(doc: DocxDocument, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

def append_suggestions_end(doc: DocxDocument, heading_text: str, content_text: str) -> tuple[Paragraph, Paragraph]:
    """Append a new section at the end: heading + content, enforcing formatting."""
    # heading
    h = doc.add_paragraph(heading_text)
    _set_paragraph_format(h)

    # two blank lines (use explicit empties; avoids stray styles from heading run)
    doc.add_paragraph("")
    doc.add_paragraph("")

    # content
    p = doc.add_paragraph(content_text)
    _set_paragraph_format(p)

    return h, p

# utils/docx_tools.py
def replace_text_in_textboxes_xml(docx_path: Path, heading_label: str, new_text: str):
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS = {"w": W_NS}
    LABELS = [heading_label, f"{heading_label}:", f"{heading_label}："]

    def para_text(p):
        return "".join((t.text or "") for t in p.xpath(".//w:t", namespaces=NS)).strip()

    def set_left_align(p):
        pPr = p.find("w:pPr", namespaces=NS)
        if pPr is None:
            pPr = etree.Element(f"{{{W_NS}}}pPr")
            p.insert(0, pPr)
        for jc in pPr.findall("w:jc", namespaces=NS):
            pPr.remove(jc)
        jc = etree.Element(f"{{{W_NS}}}jc")
        jc.set(f"{{{W_NS}}}val", "left")
        pPr.append(jc)

    def first_run(p):
        r = p.find(".//w:r", namespaces=NS)
        if r is None:
            r = etree.Element(f"{{{W_NS}}}r")
            p.append(r)
        return r

    def copy_rPr(src_r, dst_r):
        for old in dst_r.findall("w:rPr", namespaces=NS):
            dst_r.remove(old)
        rPr = src_r.find("w:rPr", namespaces=NS)
        if rPr is not None:
            dst_r.insert(0, etree.fromstring(etree.tostring(rPr)))

    def clear_all_text_after_label_in_same_para(p, keep_label_text):
        """把 label 段內，label 後面的所有字清空，但不刪節點。"""
        seen_label = False
        acc = ""
        for r in p.findall(".//w:r", namespaces=NS):
            t = r.find("w:t", namespaces=NS)
            txt = (t.text or "") if t is not None else ""
            acc += txt
            if not seen_label:
                # 尚未到達完整 label
                if keep_label_text and acc.startswith(keep_label_text) and len(acc) >= len(keep_label_text):
                    # 調整第一個 run 讓它只保留 label；其餘 run 清空
                    if t is None:
                        t = etree.Element(f"{{{W_NS}}}t")
                        r.append(t)
                    # 讓當前 run 含完整 label，之後的 run 全清空
                    t.text = keep_label_text
                    seen_label = True
                else:
                    # 還沒湊滿 label → 先清空（稍後我們會補回 label 到第一個 run）
                    if t is not None:
                        t.text = ""
            else:
                if t is not None:
                    t.text = ""

        # 若沒有任何 run 裡湊滿 label，保底：把第一個 run 設為 label
        if not seen_label:
            r0 = first_run(p)
            t0 = r0.find("w:t", namespaces=NS)
            if t0 is None:
                t0 = etree.Element(f"{{{W_NS}}}t")
                r0.append(t0)
            t0.text = keep_label_text
            # 其他 run 清空
            for r in p.findall(".//w:r", namespaces=NS):
                if r is r0:
                    continue
                t = r.find("w:t", namespaces=NS)
                if t is not None:
                    t.text = ""

    def append_run_after_label_with_style(p, text, style_src_run):
        # 追加一個 run 放「空格 + 新文字」，並複製 label 的 rPr（含底線）
        r = etree.Element(f"{{{W_NS}}}r")
        if style_src_run is not None:
            copy_rPr(style_src_run, r)
        t = etree.Element(f"{{{W_NS}}}t")
        t.text = " " + text  # label 與內文之間留一個空格
        r.append(t)
        p.append(r)

    def find_label_paragraph(root):
        # 先找 textbox 內的，找不到再找一般段
        tx_list = root.xpath(".//w:txbxContent", namespaces=NS)
        if tx_list:
            for tx in tx_list:
                for p in tx.xpath(".//w:p", namespaces=NS):
                    if any(para_text(p).startswith(lb) for lb in LABELS):
                        return p
        for p in root.xpath(".//w:p", namespaces=NS):
            if any(para_text(p).startswith(lb) for lb in LABELS):
                return p
        return None

    with zipfile.ZipFile(docx_path, "r") as zin:
        if "word/document.xml" not in zin.namelist():
            return False, None
        root = etree.fromstring(zin.read("word/document.xml"))

        label_p = find_label_paragraph(root)
        if label_p is None:
            return False, None

        # 取得實際 label 形式（含全形/半形冒號）
        full_txt = para_text(label_p)
        keep_label = next((lb for lb in LABELS if full_txt.startswith(lb)), LABELS[0])

        # 來源樣式：取 label 段第一個 run，通常包含底線
        style_src = label_p.find(".//w:r", namespaces=NS)

        # 置左＋清掉 label 後面的文字（保持同一段）
        set_left_align(label_p)
        clear_all_text_after_label_in_same_para(label_p, keep_label)

        # 在 label 後追加新文字 run（沿用底線樣式）
        append_run_after_label_with_style(label_p, new_text, style_src)

        # 打包回 docx
        buf = BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = etree.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    return True, buf.getvalue()

def enforce_whole_doc_style(doc: DocxDocument) -> None:
    """Best-effort: iterate through paragraphs and apply required fonts/spacing."""
    for p in doc.paragraphs:
        _set_paragraph_format(p)
